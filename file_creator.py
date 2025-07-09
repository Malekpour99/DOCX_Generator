import os
import re

from docx import Document
import pandas as pd


def read_csv_file(csv_path):
    """Read CSV file and return DataFrame"""
    try:
        # Read CSV with proper encoding for Persian/Farsi text
        # Use dtype=str to preserve phone numbers with leading zeros
        df = pd.read_csv(csv_path, encoding="utf-8", dtype=str)
        print(f"Successfully loaded {len(df)} records from CSV")
        return df
    except Exception as e:
        print(f"Error reading CSV file: {e}")
        return None


def read_template_docx(template_path):
    """Read the template DOCX file"""
    try:
        doc = Document(template_path)
        print("Template DOCX file loaded successfully")
        return doc
    except Exception as e:
        print(f"Error reading template DOCX file: {e}")
        return None


def replace_text_in_paragraph(paragraph, replacements):
    """Replace text in a paragraph while preserving formatting"""
    for run in paragraph.runs:
        for placeholder, replacement in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)


def replace_text_in_table(table, replacements):
    """Replace text in table cells while preserving structure"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)


def find_and_replace_in_document(doc, replacements):
    """Find and replace text throughout the document"""
    # Replace in regular paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # Replace in tables
    for table in doc.tables:
        replace_text_in_table(table, replacements)

    # Replace in headers and footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            for table in section.header.tables:
                replace_text_in_table(table, replacements)

        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            for table in section.footer.tables:
                replace_text_in_table(table, replacements)


def create_docx_from_template(template_doc, row_data, output_path):
    """Create a DOCX file from template and row data while preserving structure"""
    try:
        # Create a copy of the template document
        # We need to save and reload to create a proper copy
        temp_path = "temp_template.docx"
        template_doc.save(temp_path)
        new_doc = Document(temp_path)

        # Clean up temporary file
        if os.path.exists(temp_path):
            os.remove(temp_path)

        # Define replacements based on CSV columns and template placeholders
        replacements = {
            "{ستون آدرس گیرنده}": str(row_data.get("آدرس گیرنده", "")),
            "{ستون تلفن گیرنده}": str(row_data.get("تلفن گیرنده", "")),
            "{ستون کد سفارش}": str(row_data.get("کد سفارش", "")),
            "{ستون نام گیرنده}": str(row_data.get("نام گیرنده", "")),
            # Also handle without brackets if they exist
            "ستون آدرس گیرنده": str(row_data.get("آدرس گیرنده", "")),
            "ستون تلفن گیرنده": str(row_data.get("تلفن گیرنده", "")),
            "ستون کد سفارش": str(row_data.get("کد سفارش", "")),
            "ستون نام گیرنده": str(row_data.get("نام گیرنده", "")),
        }

        # Find and replace text in the document
        find_and_replace_in_document(new_doc, replacements)

        # Save the new document
        new_doc.save(output_path)
        return True

    except Exception as e:
        print(f"Error creating DOCX file {output_path}: {e}")
        return False


def generate_filename(row_data, index):
    """Generate filename based on row data"""
    # Use order code and customer name for filename
    order_code = str(row_data.get("کد سفارش", f"order_{index}")).replace("/", "_")
    customer_name = str(row_data.get("نام گیرنده", f"customer_{index}")).replace(
        " ", "_"
    )

    # Clean filename from invalid characters
    filename = f"{order_code}_{customer_name}.docx"
    # Remove invalid characters for filename
    filename = re.sub(r'[<>:"/\\|?*]', "_", filename)

    return filename


def analyze_template_structure(template_doc):
    """Analyze the template structure and show placeholders"""
    print("=== Template Structure Analysis ===")

    # Analyze paragraphs
    print("\nParagraphs:")
    for i, paragraph in enumerate(template_doc.paragraphs):
        if paragraph.text.strip():
            print(f"  {i+1}. {paragraph.text[:100]}...")

    # Analyze tables
    print(f"\nTables found: {len(template_doc.tables)}")
    for i, table in enumerate(template_doc.tables):
        print(f"\n  Table {i+1}:")
        print(f"    Rows: {len(table.rows)}, Columns: {len(table.columns)}")
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"    Row {row_idx+1}, Col {col_idx+1}: {cell.text[:50]}...")

    # Find placeholders
    print("\n=== Placeholders Found ===")
    placeholders = set()

    # Check paragraphs
    for paragraph in template_doc.paragraphs:
        if "{" in paragraph.text and "}" in paragraph.text:
            # Extract placeholders
            found_placeholders = re.findall(r"\{[^}]+\}", paragraph.text)
            placeholders.update(found_placeholders)

    # Check tables
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{" in cell.text and "}" in cell.text:
                    found_placeholders = re.findall(r"\{[^}]+\}", cell.text)
                    placeholders.update(found_placeholders)

    if placeholders:
        print("Found placeholders:")
        for placeholder in sorted(placeholders):
            print(f"  - {placeholder}")
    else:
        print("No placeholders found with {} format")

    print("=" * 50)


def main():
    """Main function to process CSV and generate DOCX files"""

    # Configuration
    csv_file_path = "csv.csv"  # Your CSV file path
    template_file_path = "Forward Template.docx"  # Your template DOCX file path
    output_directory = "generated_documents"  # Output directory for DOCX files

    # Create output directory if it doesn't exist
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
        print(f"Created output directory: {output_directory}")

    # Read template document
    template_doc = read_template_docx(template_file_path)
    if template_doc is None:
        print(f"Error: Could not read template file '{template_file_path}'")
        print("Make sure the file exists and is a valid DOCX file.")
        return

    # Analyze template structure
    analyze_template_structure(template_doc)

    # Read CSV data
    df = read_csv_file(csv_file_path)
    if df is None:
        return

    print(f"\nProcessing {len(df)} records...")

    # Process each row
    successful_count = 0
    failed_count = 0

    for index, row in df.iterrows():
        try:
            # Generate filename
            filename = generate_filename(row, index + 1)
            output_path = os.path.join(output_directory, filename)

            # Create DOCX file from template
            if create_docx_from_template(template_doc, row, output_path):
                successful_count += 1
                print(f"✓ Created: {filename}")
            else:
                failed_count += 1
                print(f"✗ Failed: {filename}")

        except Exception as e:
            failed_count += 1
            print(f"✗ Error processing row {index + 1}: {e}")

    # Print summary
    print(f"\n=== Processing Summary ===")
    print(f"Total records: {len(df)}")
    print(f"Successfully processed: {successful_count}")
    print(f"Failed: {failed_count}")
    print(f"Output directory: {output_directory}")
    print(f"Files ready for use!")


def check_files_exist():
    """Check if required files exist"""
    required_files = ["csv.csv", "Forward Template.docx"]
    missing_files = []

    for file_path in required_files:
        if not os.path.exists(file_path):
            missing_files.append(file_path)

    if missing_files:
        print("Missing required files:")
        for file_path in missing_files:
            print(f"  - {file_path}")
        print("\nMake sure these files are in the same directory as the script.")
        return False

    return True


if __name__ == "__main__":
    # Check if files exist
    if not check_files_exist():
        exit(1)

    main()
